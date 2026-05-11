from __future__ import annotations

import logging
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from .models import DocumentText
from .ocr import TesseractOcr
from .utils import compact_text

logger = logging.getLogger(__name__)


class DocumentExtractor:
    def __init__(self, ocr: TesseractOcr, min_pdf_text_chars: int = 400, ocr_enabled: bool = False) -> None:
        self.ocr = ocr
        self.min_pdf_text_chars = min_pdf_text_chars
        self.ocr_enabled = ocr_enabled

    def extract_many(self, paths: list[Path]) -> list[DocumentText]:
        result: list[DocumentText] = []
        for path in paths:
            result.append(self.extract(path))
        return result

    def extract(self, path: Path) -> DocumentText:
        suffix = path.suffix.lower()
        logger.info("Extracting document: %s", path)
        if suffix == ".pdf":
            return self._extract_pdf(path)
        if suffix in {".docx", ".doc"}:
            return self._extract_word(path)
        if suffix in {".xlsx", ".xls"}:
            return self._extract_excel(path)
        raise ValueError(f"Unsupported file format: {path.suffix}. Use PDF, DOCX, XLSX or DOC.")

    def _extract_word(self, path: Path) -> DocumentText:
        if path.suffix.lower() == ".doc":
            raise ValueError("Legacy .doc is not supported directly. Convert it to .docx or PDF first.")
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is not installed. Install dependencies from requirements.txt.") from exc

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    tables.append(" | ".join(cells))
        text = compact_text("\n".join(paragraphs + tables))
        return DocumentText(path=path, kind="word", text=text, used_ocr=False)

    def _extract_excel(self, path: Path) -> DocumentText:
        if path.suffix.lower() == ".xls":
            raise ValueError("Legacy .xls is not supported directly. Convert it to .xlsx first.")

        workbook = self._read_xlsx_text(path)
        text = compact_text("\n".join(workbook))
        return DocumentText(path=path, kind="excel", text=text, used_ocr=False)

    def _read_xlsx_text(self, path: Path) -> list[str]:
        parts: list[str] = []
        with zipfile.ZipFile(path) as archive:
            shared_strings = self._read_shared_strings(archive)
            sheets = self._read_workbook_sheets(archive)
            for sheet_name, sheet_path in sheets:
                if sheet_path not in archive.namelist():
                    continue
                parts.append(f"### Лист: {sheet_name}")
                for row in self._iter_xlsx_rows(archive, sheet_path, shared_strings):
                    values = [self._cell_to_text(value) for value in row]
                    while values and not values[-1]:
                        values.pop()
                    if any(values):
                        parts.append(" | ".join(values))
        return parts

    @staticmethod
    def _read_shared_strings(archive: zipfile.ZipFile) -> list[str]:
        if "xl/sharedStrings.xml" not in archive.namelist():
            return []
        root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
        ns = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        strings: list[str] = []
        for item in root.findall("x:si", ns):
            texts = [node.text or "" for node in item.findall(".//x:t", ns)]
            strings.append("".join(texts))
        return strings

    @staticmethod
    def _read_workbook_sheets(archive: zipfile.ZipFile) -> list[tuple[str, str]]:
        ns = {
            "x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
        }
        workbook = ET.fromstring(archive.read("xl/workbook.xml"))
        rels = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
        rel_targets = {
            rel.attrib["Id"]: rel.attrib["Target"]
            for rel in rels.findall("pr:Relationship", ns)
        }

        sheets: list[tuple[str, str]] = []
        for sheet in workbook.findall("x:sheets/x:sheet", ns):
            name = sheet.attrib.get("name", "Sheet")
            rel_id = sheet.attrib.get(f"{{{ns['r']}}}id")
            target = rel_targets.get(rel_id or "")
            if not target:
                continue
            if target.startswith("/"):
                normalized = target.lstrip("/")
            elif target.startswith("xl/"):
                normalized = target
            else:
                normalized = f"xl/{target}"
            sheets.append((name, normalized.replace("\\", "/")))
        return sheets

    @staticmethod
    def _iter_xlsx_rows(
        archive: zipfile.ZipFile,
        sheet_path: str,
        shared_strings: list[str],
    ) -> list[list[str]]:
        root = ET.fromstring(archive.read(sheet_path))
        ns = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        rows: list[list[str]] = []
        for row in root.findall(".//x:sheetData/x:row", ns):
            values: list[str] = []
            for cell in row.findall("x:c", ns):
                cell_type = cell.attrib.get("t")
                if cell_type == "inlineStr":
                    texts = [node.text or "" for node in cell.findall(".//x:t", ns)]
                    values.append("".join(texts))
                    continue
                value_node = cell.find("x:v", ns)
                if value_node is None or value_node.text is None:
                    values.append("")
                    continue
                raw = value_node.text
                if cell_type == "s":
                    try:
                        values.append(shared_strings[int(raw)])
                    except (ValueError, IndexError):
                        values.append(raw)
                elif cell_type == "b":
                    values.append("Да" if raw == "1" else "Нет")
                else:
                    values.append(raw)
            while values and not values[-1]:
                values.pop()
            if any(values):
                rows.append(values)
        return rows

    @staticmethod
    def _cell_to_text(value: object) -> str:
        if value is None:
            return ""
        return str(value).strip()

    def _extract_pdf(self, path: Path) -> DocumentText:
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError("PyMuPDF is not installed. Install dependencies from requirements.txt.") from exc

        with fitz.open(path) as pdf:
            text = compact_text("\n".join(page.get_text("text") for page in pdf))
        if len(text) >= self.min_pdf_text_chars:
            return DocumentText(path=path, kind="pdf", text=text, used_ocr=False)

        if not self.ocr_enabled:
            if text:
                logger.warning("PDF text layer is short and OCR is disabled: %s", path)
                return DocumentText(path=path, kind="pdf", text=text, used_ocr=False)
            raise RuntimeError(
                "PDF похож на скан: текстовый слой не найден, а OCR отключен. "
                "Загрузите DOCX/PDF с распознаваемым текстом или включите OCR после установки Tesseract "
                "и Python-зависимостей."
            )

        logger.info("PDF text layer is short (%s chars), running OCR: %s", len(text), path)
        try:
            ocr_text = self.ocr.recognize_pdf(path)
        except RuntimeError:
            if text:
                logger.exception("OCR failed; using short PDF text layer: %s", path)
                return DocumentText(path=path, kind="pdf", text=text, used_ocr=False)
            raise
        merged = compact_text(text + "\n\n" + ocr_text)
        return DocumentText(path=path, kind="pdf", text=merged, used_ocr=True)
