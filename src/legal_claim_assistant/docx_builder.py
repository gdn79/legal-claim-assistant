from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


CHAT_TAIL_PATTERNS = (
    re.compile(r"^\s*если\s+хотите\b", re.IGNORECASE),
    re.compile(r"^\s*если\s+необходимо\b", re.IGNORECASE),
    re.compile(r"^\s*при\s+необходимости\b", re.IGNORECASE),
    re.compile(r"^\s*хотите,\s+чтобы\s+я\b", re.IGNORECASE),
    re.compile(r"^\s*могу\s+подготовить\b", re.IGNORECASE),
    re.compile(r"^\s*я\s+могу\s+подготовить\b", re.IGNORECASE),
    re.compile(r"^\s*готов\s+подготовить\b", re.IGNORECASE),
)
AI_NOTE_HEADING_PATTERNS = (
    re.compile(r"^\s*\**\s*примечани[ея]\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*\**\s*важные\s+уточнения\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*\**\s*дополнительные\s+уточнения\s*:?\s*\**\s*$", re.IGNORECASE),
)
SOURCE_BLOCK_PATTERNS = (
    re.compile(r"^\s*\**\s*ссылки\s+на\s+(?:источники|нормы|правовые\s+источники)\b.*$", re.IGNORECASE),
    re.compile(r"^\s*\**\s*источники\s+(?:права|норм|информации)\b.*$", re.IGNORECASE),
    re.compile(r"^\s*\**\s*нормативные\s+источники\b.*$", re.IGNORECASE),
    re.compile(r"^\s*\**\s*использованные\s+источники\b.*$", re.IGNORECASE),
)
SOURCE_LINE_PATTERNS = (
    re.compile(r"https?://", re.IGNORECASE),
    re.compile(r"\bconsultant\.ru\b", re.IGNORECASE),
    re.compile(r"\bgarant\.ru\b", re.IGNORECASE),
    re.compile(r"\bpravo\.gov\.ru\b", re.IGNORECASE),
    re.compile(r"^\s*(?:ГК|АПК|НК)\s+РФ\s*[—-]\s*https?://", re.IGNORECASE),
)
SECTION_HEADING_PATTERNS = (
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*исковое\s+заявление\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*истец\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*ответчик\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*шапка\s*(?:иска|заявления)?\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*цена\s+иска\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*госпошлина\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*суть\s+(?:спора|требований)\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*обстоятельства\s+(?:дела|спора)\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*правов(?:ое|ые)\s+(?:основание|основания|обоснование)\s*(?:требований)?\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*расч[её]т\s*(?:задолженности|требований)?\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*сумма\s+иска\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*претензионн(?:ый|ого)\s+поряд(?:ок|ка)\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*просьба\s+истца\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*просительн(?:ая|ую)\s+част[ьи]\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*приложени[ея]\s*:?\s*\**\s*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*дата\s+подачи\s+иска\s*:?\s*\**.*$", re.IGNORECASE),
    re.compile(r"^\s*(?:\d+[.)]\s*)?\**\s*представитель\s+истца\s*:?\s*\**\s*$", re.IGNORECASE),
)


def clean_claim_text(claim_text: str) -> str:
    cleaned_lines: list[str] = []
    blank_pending = False
    for raw_line in claim_text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.strip()
        if line in {"---", "***", "___"}:
            continue
        if line.startswith("```"):
            continue
        if any(pattern.search(line) for pattern in AI_NOTE_HEADING_PATTERNS):
            break
        if any(pattern.search(line) for pattern in CHAT_TAIL_PATTERNS):
            break
        if any(pattern.search(line) for pattern in SOURCE_BLOCK_PATTERNS):
            break
        if any(pattern.search(line) for pattern in SOURCE_LINE_PATTERNS):
            continue
        if any(pattern.search(line) for pattern in SECTION_HEADING_PATTERNS):
            continue
        line = _clean_inline_markdown(line)
        if not line:
            blank_pending = bool(cleaned_lines)
            continue
        if blank_pending and cleaned_lines and cleaned_lines[-1].strip():
            cleaned_lines.append("")
        blank_pending = False
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()


def _clean_inline_markdown(line: str) -> str:
    line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
    line = re.sub(r"__(.*?)__", r"\1", line)
    return line.strip()


def build_claim_docx(claim_text: str, output_path: Path) -> Path:
    claim_text = clean_claim_text(claim_text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ИСКОВОЕ ЗАЯВЛЕНИЕ")
    run.bold = True
    run.font.size = Pt(14)

    for block in claim_text.split("\n"):
        line = block.strip()
        if not line:
            doc.add_paragraph()
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(28)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.add_run(line)

    doc.save(output_path)
    return output_path
